from docx import Document
import logging

logger = logging.getLogger(__name__)

class WordNode:
    def __init__(self, config):
        self.config = config
        logger.info(f"Initialized WordNode with config: {self.config}")

    def execute(self, input_data=None):
        try:
            logger.info(f"Word node executing with input data: {input_data}")
            
            filename = self.config.get('filename')
            content = self.config.get('content', '')
            
            logger.info(f"Creating Word document: {filename}")
            logger.info(f"Content: {content}")

            doc = Document()
            
            # 添加标题
            doc.add_heading('数据处理报告', 0)
            
            # 添加基本内容
            doc.add_paragraph(content)

            # 添加Excel数据
            if input_data:
                doc.add_heading('数据统计结果', level=1)
                for source_id, source_data in input_data.items():
                    logger.info(f"Processing data from {source_id}: {source_data}")
                    if isinstance(source_data, dict):
                        doc.add_heading(f'节点 {source_id} 的数据', level=2)
                        if 'data' in source_data:
                            for k, v in source_data['data'].items():
                                doc.add_paragraph(f"{k}: {v}")
                        elif 'status' in source_data and source_data['status'] == 'success':
                            for k, v in source_data.items():
                                if k != 'status':
                                    doc.add_paragraph(f"{k}: {v}")

            logger.info(f"Saving document to: {filename}")
            doc.save(filename)
            logger.info("Document saved successfully")

            return {
                "status": "success",
                "message": f"文档已保存到: {filename}"
            }

        except Exception as e:
            logger.error(f"Error in Word node: {str(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "status": "error",
                "message": str(e)
            }