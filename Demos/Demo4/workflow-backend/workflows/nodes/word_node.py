from .base_node import BaseNode
from docx import Document
import logging

logger = logging.getLogger(__name__)

class WordNode(BaseNode):
    def __init__(self, id: str, next_data: dict, node_data: dict):
        super().__init__(id, "word", next_data, node_data)
        self.config = node_data.get('data', {})
        
        if len(self._next_data["out"]) > 0:
            self._next_node = self._next_data["out"][0][1]
            self._out_handle = self._next_data["out"][0][2]
        
        if len(self._next_data["in"]) > 0:
            self._in_handle = self._next_data["in"][0][2]
            
        logger.info(f"初始化Word节点: {id}")

    def run(self):
        try:
            logger.info(f"Word节点执行，输入数据: {self.input_data}")
            
            filename = self.config.get('filename')
            content = self.config.get('content', '')
            
            if not filename:
                raise ValueError("需要指定Word文件路径")
            
            logger.info(f"创建Word文档: {filename}")
            logger.info(f"内容: {content}")

            doc = Document()
            
            # 添加标题
            doc.add_heading('数据处理报告', 0)
            
            # 添加基本内容
            doc.add_paragraph(content)

            # 添加输入数据
            if self.input_data:
                doc.add_heading('数据统计结果', level=1)
                if isinstance(self.input_data, dict):
                    for k, v in self.input_data.items():
                        doc.add_paragraph(f"{k}: {v}")
                elif isinstance(self.input_data, list):
                    for item in self.input_data:
                        if isinstance(item, dict):
                            for k, v in item.items():
                                doc.add_paragraph(f"{k}: {v}")
                        else:
                            doc.add_paragraph(str(item))
                else:
                    doc.add_paragraph(str(self.input_data))

            logger.info(f"保存文档到: {filename}")
            doc.save(filename)
            logger.info("文档保存成功")

            self.output_data = filename
            return f"Word文档已保存到: {filename}"

        except Exception as e:
            logger.error(f"Word节点错误: {str(e)}")
            import traceback
            logger.error(f"错误追踪: {traceback.format_exc()}")
            raise ValueError(str(e))